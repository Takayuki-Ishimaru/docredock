#!/usr/bin/env python3
"""Generate the deterministic, element-rich DOCX corpus used by DRMD DOCX tests.

Companion to tests/DocRedock.Tests/Fixtures/Pdf/generate_complex_pdf.py. Builds a Word
design document that exercises D01-D18 from COMPLEX_DESIGN_DOC_SPEC.md, reusing the
same facts (IF-05, TR-09, TX-01, BR-01, F-06, API-01, TC-009, NFR-AVL-02, ISSUE-03,
outbox_event, ...) as the source-of-truth Markdown projection of
経費精算システム_設計書_検証用.xlsx, and embedding IMG-01/IMG-02 extracted from that
same workbook.

Usage:
  python3 generate_complex_docx.py [output.docx] [--xlsx path] [--root path]

Where python-docx 1.2.0 has no write API (TOC field, PAGE field, footnotes,
endnotes, hyperlinks, bookmarks, tracked changes, text boxes, floating images),
OOXML is injected directly via lxml against the python-docx object model.
"""
from __future__ import annotations

import argparse
import io
import sys
import zipfile
from datetime import datetime
from pathlib import Path

from lxml import etree

import docx
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

# --------------------------------------------------------------------------
# Paths and constants
# --------------------------------------------------------------------------

HERE = Path(__file__).resolve().parent
DEFAULT_ROOT = HERE.parents[3] if len(HERE.parents) >= 3 else HERE
DEFAULT_XLSX = DEFAULT_ROOT / "経費精算システム_設計書_検証用.xlsx"
DEFAULT_OUT = HERE / "complex-design-doc.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

NAVY = "17365D"
TEAL = "0B7285"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E6F7F8"
MUTED = "667085"
RED_HEX = "B42318"

FIXED_DATETIME = datetime(2026, 8, 23, 9, 0, 0)
REVIEW_AUTHOR = "レビュー担当"
REVIEW_DATE_ISO = "2026-08-24T10:15:00Z"

CHECKLIST: list[tuple[str, str, str, str]] = []  # (id, label, method, note)

def check(item_id: str, label: str, method: str, note: str = "") -> None:
    CHECKLIST.append((item_id, label, method, note))

# --------------------------------------------------------------------------
# Low-level OOXML helpers
# --------------------------------------------------------------------------

def frag(xml: str) -> etree._Element:
    """Parse a standalone XML fragment (with its own xmlns declarations)."""
    return etree.fromstring(xml.encode("utf-8"))

def append_raw(paragraph, xml: str) -> etree._Element:
    element = frag(xml)
    paragraph._p.append(element)
    return element

def replace_paragraph_xml(paragraph, xml: str) -> etree._Element:
    new_p = frag(xml)
    paragraph._p.getparent().replace(paragraph._p, new_p)
    return new_p

def shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_bold(cell, bold: bool = True) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold

def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def header_row(table, values, fill=NAVY):
    for index, value in enumerate(values):
        cell = table.cell(0, index)
        set_cell_text(cell, value, bold=True, color="FFFFFF")
        shade_cell(cell, fill)

def data_row(table, row_index, values):
    for index, value in enumerate(values):
        set_cell_text(table.cell(row_index, index), value)

def build_table(document, headers, rows, style="Table Grid"):
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style
    header_row(table, headers)
    for offset, row in enumerate(rows, start=1):
        data_row(table, offset, row)
    return table

# --------------------------------------------------------------------------
# Hyperlinks and bookmarks (no python-docx write API in 1.2.0)
# --------------------------------------------------------------------------

_BOOKMARK_ID = [100]

def add_external_hyperlink(paragraph, text: str, url: str) -> None:
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    xml = f'''<w:hyperlink xmlns:w="{W}" xmlns:r="{R}" r:id="{rid}" w:history="1">
      <w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>
    </w:hyperlink>'''
    append_raw(paragraph, xml)

def add_internal_hyperlink(paragraph, text: str, bookmark_name: str) -> None:
    xml = f'''<w:hyperlink xmlns:w="{W}" w:anchor="{bookmark_name}" w:history="1">
      <w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>
    </w:hyperlink>'''
    append_raw(paragraph, xml)

def add_bookmark(paragraph, name: str) -> None:
    _BOOKMARK_ID[0] += 1
    bid = _BOOKMARK_ID[0]
    append_raw(paragraph, f'<w:bookmarkStart xmlns:w="{W}" w:id="{bid}" w:name="{name}"/>')
    append_raw(paragraph, f'<w:bookmarkEnd xmlns:w="{W}" w:id="{bid}"/>')

# --------------------------------------------------------------------------
# Multi-level list numbering with EXPLICIT per-paragraph ilvl.
#
# python-docx's default template gives "List Bullet"/"List Bullet 2"/"List
# Bullet 3" each their OWN numId with only ilvl=0 defined (singleLevel) — the
# indent differs only because each style points at a different numId, never
# because a paragraph's own w:numPr/w:ilvl varies. DRMD's DocxAdapter.ListLevel
# reads ONLY the paragraph's own w:numPr/w:ilvl, so those default styles would
# report level 0 for every item and defeat the D09 nesting check. Register one
# real multi-level abstractNum instead and stamp w:ilvl explicitly per item,
# matching how Word itself encodes a Tab-indented nested bullet list.
# --------------------------------------------------------------------------

def add_multilevel_bullet_numbering(document) -> int:
    numbering_root = document.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering_root.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering_root.findall(qn("w:num"))]
    new_abstract_id = max(abstract_ids) + 1
    new_num_id = max(num_ids) + 1

    bullet_chars = ["", "o", ""]  # Symbol/Courier New/Wingdings bullets, matching Word's own L1-L3 convention
    bullet_fonts = ["Symbol", "Courier New", "Wingdings"]

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(new_abstract_id))
    multilevel = OxmlElement("w:multiLevelType")
    multilevel.set(qn("w:val"), "hybridMultilevel")
    abstract_num.append(multilevel)
    for ilvl in range(3):
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_format = OxmlElement("w:numFmt")
        num_format.set(qn("w:val"), "bullet")
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), bullet_chars[ilvl])
        justify = OxmlElement("w:lvlJc")
        justify.set(qn("w:val"), "left")
        level_pPr = OxmlElement("w:pPr")
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(360 + ilvl * 360))
        indent.set(qn("w:hanging"), "360")
        level_pPr.append(indent)
        level_rPr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), bullet_fonts[ilvl])
        fonts.set(qn("w:hAnsi"), bullet_fonts[ilvl])
        fonts.set(qn("w:hint"), "default")
        level_rPr.append(fonts)
        for child in (start, num_format, level_text, justify, level_pPr, level_rPr):
            level.append(child)
        abstract_num.append(level)

    numbering_root.findall(qn("w:abstractNum"))[-1].addnext(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(new_abstract_id))
    num.append(abstract_ref)
    numbering_root.append(num)

    return new_num_id

def set_explicit_list_level(paragraph, num_id: int, ilvl: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    pPr.append(num_pr)

# --------------------------------------------------------------------------
# Tracked changes (w:ins / w:del)
# --------------------------------------------------------------------------

_REVISION_ID = [200]

def add_tracked_insertion(paragraph, text: str) -> None:
    _REVISION_ID[0] += 1
    xml = (
        f'<w:ins xmlns:w="{W}" w:id="{_REVISION_ID[0]}" w:author="{REVIEW_AUTHOR}" '
        f'w:date="{REVIEW_DATE_ISO}"><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:ins>'
    )
    append_raw(paragraph, xml)

def add_tracked_deletion(paragraph, text: str) -> None:
    _REVISION_ID[0] += 1
    xml = (
        f'<w:del xmlns:w="{W}" w:id="{_REVISION_ID[0]}" w:author="{REVIEW_AUTHOR}" '
        f'w:date="{REVIEW_DATE_ISO}"><w:r><w:delText xml:space="preserve">{text}</w:delText></w:r></w:del>'
    )
    append_raw(paragraph, xml)

# --------------------------------------------------------------------------
# Text box (legacy VML: still the simplest reliably-openable form)
# --------------------------------------------------------------------------

_SHAPE_ID = [1024]

def add_textbox(paragraph, text: str, width_pt=300, height_pt=54) -> None:
    _SHAPE_ID[0] += 1
    shape_id = f"_x0000_s{_SHAPE_ID[0]}"
    xml = f'''<w:r xmlns:w="{W}" xmlns:v="urn:schemas-microsoft-com:vml"
        xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word">
      <w:pict>
        <v:shape id="TextBox{_SHAPE_ID[0]}" o:spid="{shape_id}" type="#_x0000_t202"
          style="position:absolute;margin-left:0;margin-top:6pt;width:{width_pt}pt;height:{height_pt}pt;z-index:2"
          fillcolor="#{LIGHT_BLUE}" strokecolor="#{NAVY}" strokeweight="1pt">
          <v:textbox inset="6pt,4pt,6pt,4pt">
            <w:txbxContent>
              <w:p><w:pPr><w:jc w:val="left"/></w:pPr>
                <w:r><w:rPr><w:b/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:shape>
      </w:pict>
    </w:r>'''
    append_raw(paragraph, xml)

# --------------------------------------------------------------------------
# Floating (anchored) picture: convert an inline picture run to wp:anchor
# --------------------------------------------------------------------------

def convert_last_picture_to_anchor(paragraph, wrap="square", offset_x=228600, offset_y=114300) -> None:
    run = paragraph.runs[-1]
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    extent = inline.find(qn("wp:extent"))
    effect_extent = inline.find(qn("wp:effectExtent"))
    doc_pr = inline.find(qn("wp:docPr"))
    frame_pr = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    for key, value in {
        "distT": "0", "distB": "0", "distL": "114300", "distR": "114300",
        "simplePos": "0", "relativeHeight": "251659264", "behindDoc": "0",
        "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
    }.items():
        anchor.set(key, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "column")
    pos_offset_h = OxmlElement("wp:posOffset")
    pos_offset_h.text = str(offset_x)
    position_h.append(pos_offset_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    pos_offset_v = OxmlElement("wp:posOffset")
    pos_offset_v.text = str(offset_y)
    position_v.append(pos_offset_v)

    wrap_square = OxmlElement(f"wp:wrap{wrap.capitalize()}")
    wrap_square.set("wrapText", "bothSides")

    anchor.append(simple_pos)
    anchor.append(position_h)
    anchor.append(position_v)
    anchor.append(extent)
    if effect_extent is not None:
        anchor.append(effect_extent)
    anchor.append(wrap_square)
    anchor.append(doc_pr)
    anchor.append(frame_pr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)

# --------------------------------------------------------------------------
# Footnotes / endnotes (new OPC parts; no python-docx write API)
# --------------------------------------------------------------------------

class NoteBook:
    """Accumulates footnote/endnote bodies and attaches the parts once at the end."""

    def __init__(self):
        self.footnotes: list[tuple[int, str]] = []
        self.endnotes: list[tuple[int, str]] = []
        self._next_footnote_id = 2
        self._next_endnote_id = 2

    def add_footnote(self, paragraph, text: str) -> int:
        note_id = self._next_footnote_id
        self._next_footnote_id += 1
        self.footnotes.append((note_id, text))
        xml = (
            f'<w:r xmlns:w="{W}"><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
            f'<w:footnoteReference w:id="{note_id}"/></w:r>'
        )
        append_raw(paragraph, xml)
        return note_id

    def add_endnote(self, paragraph, text: str) -> int:
        note_id = self._next_endnote_id
        self._next_endnote_id += 1
        self.endnotes.append((note_id, text))
        xml = (
            f'<w:r xmlns:w="{W}"><w:rPr><w:rStyle w:val="EndnoteReference"/></w:rPr>'
            f'<w:endnoteReference w:id="{note_id}"/></w:r>'
        )
        append_raw(paragraph, xml)
        return note_id

    def attach(self, document) -> None:
        footnotes_xml = self._build_notes_part(
            root_tag="w:footnotes", entry_tag="w:footnote", ref_tag="w:footnoteRef",
            style_id="FootnoteText", entries=self.footnotes,
        )
        endnotes_xml = self._build_notes_part(
            root_tag="w:endnotes", entry_tag="w:endnote", ref_tag="w:endnoteRef",
            style_id="EndnoteText", entries=self.endnotes,
        )
        footnotes_part = Part(PackURI("/word/footnotes.xml"), CT.WML_FOOTNOTES, footnotes_xml, document.part.package)
        endnotes_part = Part(PackURI("/word/endnotes.xml"), CT.WML_ENDNOTES, endnotes_xml, document.part.package)
        document.part.relate_to(footnotes_part, RT.FOOTNOTES)
        document.part.relate_to(endnotes_part, RT.ENDNOTES)

    @staticmethod
    def _build_notes_part(*, root_tag, entry_tag, ref_tag, style_id, entries) -> bytes:
        short = root_tag.split(":")[1][:-1]  # "footnotes" -> "footnote", "endnotes" -> "endnote"
        parts = [
            f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
            f'<{root_tag} xmlns:w="{W}">',
            f'<w:{short} w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:{short}>',
            f'<w:{short} w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p></w:{short}>',
        ]
        for note_id, text in entries:
            parts.append(
                f'<w:{short} w:id="{note_id}"><w:p><w:pPr><w:pStyle w:val="{style_id}"/></w:pPr>'
                f'<w:r><w:rPr><w:rStyle w:val="{ref_tag[2].upper()}{ref_tag[3:]}"/></w:rPr><w:{ref_tag[2:]}/></w:r>'
                f'<w:r><w:t xml:space="preserve"> {text}</w:t></w:r></w:p></w:{short}>'
            )
        parts.append(f"</{root_tag}>")
        return "".join(parts).encode("utf-8")

# --------------------------------------------------------------------------
# TOC field (complex field: begin / instrText / separate ... end)
# --------------------------------------------------------------------------

def insert_toc(document, entries: list[tuple[int, str, int]]) -> None:
    """entries: list of (level, title, cached_page_number)."""
    begin = document.add_paragraph()
    replace_paragraph_xml(begin, f'''<w:p xmlns:w="{W}">
      <w:pPr><w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9350"/></w:tabs></w:pPr>
      <w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>
      <w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
    </w:p>''')

    for level, title, page in entries:
        indent = 220 * max(0, level - 1)
        entry = document.add_paragraph()
        replace_paragraph_xml(entry, f'''<w:p xmlns:w="{W}">
          <w:pPr>
            <w:ind w:left="{indent}"/>
            <w:tabs><w:tab w:val="right" w:leader="dot" w:pos="9350"/></w:tabs>
          </w:pPr>
          <w:r><w:t xml:space="preserve">{title}</w:t></w:r>
          <w:r><w:tab/></w:r>
          <w:r><w:t xml:space="preserve">{page}</w:t></w:r>
        </w:p>''')

    end = document.add_paragraph()
    replace_paragraph_xml(end, f'<w:p xmlns:w="{W}"><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>')

# --------------------------------------------------------------------------
# PAGE field in the footer (simple field is enough; no separate/end needed)
# --------------------------------------------------------------------------

def add_page_field_footer(section, prefix: str) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.add_run(prefix)
    xml = f'''<w:fldSimple xmlns:w="{W}" w:instr="PAGE">
      <w:r><w:t>1</w:t></w:r>
    </w:fldSimple>'''
    append_raw(paragraph, xml)

# --------------------------------------------------------------------------
# Styles that Word's default template does not ship (footnotes/endnotes/TOC/
# hyperlink/code all need a real style definition to stay Word-safe).
# --------------------------------------------------------------------------

def ensure_supporting_styles(document) -> None:
    styles = document.styles

    hyperlink = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    hyperlink.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    hyperlink.font.underline = True

    footnote_ref = styles.add_style("FootnoteReference", WD_STYLE_TYPE.CHARACTER)
    footnote_ref.font.superscript = True

    endnote_ref = styles.add_style("EndnoteReference", WD_STYLE_TYPE.CHARACTER)
    endnote_ref.font.superscript = True

    footnote_text = styles.add_style("FootnoteText", WD_STYLE_TYPE.PARAGRAPH)
    footnote_text.base_style = styles["Normal"]
    footnote_text.font.size = Pt(9)

    endnote_text = styles.add_style("EndnoteText", WD_STYLE_TYPE.PARAGRAPH)
    endnote_text.base_style = styles["Normal"]
    endnote_text.font.size = Pt(9)

    source_code = styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
    source_code.base_style = styles["Normal"]
    source_code.font.name = "Consolas"
    source_code.font.size = Pt(9.5)
    source_code.paragraph_format.space_before = Pt(4)
    source_code.paragraph_format.space_after = Pt(8)
    pPr = source_code.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F7FA")
    pPr.append(shd)

def set_east_asian_font(run, name="Yu Gothic") -> None:
    rpr = run._r.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), name)

# --------------------------------------------------------------------------
# Image extraction from the sibling XLSX corpus (IMG-01 / IMG-02)
# --------------------------------------------------------------------------

def extract_images(xlsx_path: Path) -> tuple[bytes, bytes]:
    with zipfile.ZipFile(xlsx_path) as archive:
        img01 = archive.read("xl/media/image.png")
        img02 = archive.read("xl/media/image2.png")
    return img01, img02

# --------------------------------------------------------------------------
# Document assembly
# --------------------------------------------------------------------------

def build(output_path: Path, xlsx_path: Path) -> Path:
    img01_bytes, img02_bytes = extract_images(xlsx_path)

    document = docx.Document()
    notes = NoteBook()
    ensure_supporting_styles(document)

    document.core_properties.title = "経費精算システム 基本・詳細設計書（検証用）"
    document.core_properties.subject = "DRMD DOCX変換検証用フィクスチャ"
    document.core_properties.author = "アプリ設計担当"
    document.core_properties.created = FIXED_DATETIME
    document.core_properties.modified = FIXED_DATETIME

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    set_east_asian_font_style(normal)

    for level, size in ((1, 16), (2, 13), (3, 11.5), (4, 10.5)):
        style = document.styles[f"Heading {level}"]
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        set_east_asian_font_style(style)

    section = document.sections[0]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "経費精算システム 基本・詳細設計書"
    add_page_field_footer(section, "ページ ")
    check("D04", "ヘッダー(文書名)/フッター(PAGEフィールド)", "python-docx section.header (native) + raw w:fldSimple PAGE (oxml)")

    build_cover(document)

    toc_entries = [
        (1, "第1章　概要", 3),
        (1, "第2章　状態遷移", 5),
        (1, "第3章　シーケンス", 7),
        (1, "第4章　業務フロー", 8),
        (1, "第5章　画面・入力仕様", 10),
        (1, "第6章　API・データ設計", 11),
        (1, "第7章　テストシナリオ", 13),
        (1, "第8章　非機能・運用設計", 14),
        (1, "付録", 16),
    ]
    insert_toc(document, toc_entries)
    check("D03", "目次(TOCフィールド、キャッシュ済みエントリ)", "raw oxml: w:fldChar begin/separate/end + w:instrText ' TOC \\\\o ' + cached entry paragraphs")

    document.add_page_break()

    build_chapter1(document)
    build_chapter2(document)
    build_chapter3(document)
    build_chapter4(document)
    build_chapter5(document, img01_bytes, img02_bytes)
    build_chapter6(document)
    build_chapter7(document)
    landscape_bookmark = build_chapter8(document)
    build_appendix(document, notes, landscape_bookmark)

    notes.attach(document)
    check("D16", "脚注1件+文末脚注1件", "raw oxml: new word/footnotes.xml / word/endnotes.xml parts via docx.opc.part.Part + relate_to, w:footnoteReference/w:endnoteReference runs")

    document.save(output_path)
    return output_path

def set_east_asian_font_style(style) -> None:
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), "Yu Gothic")

def build_cover(document) -> None:
    title = document.add_paragraph()
    title.style = document.styles["Title"]
    run = title.add_run("経費精算システム")
    set_east_asian_font(run)
    title2 = document.add_paragraph()
    title2.style = document.styles["Title"]
    run2 = title2.add_run("基本・詳細設計書（検証用サンプル）")
    set_east_asian_font(run2)

    meta_rows = [
        ("文書番号", "EXPS-DES-001"),
        ("版", "1.2"),
        ("状態", "レビュー中（検証用）"),
        ("作成日", "2026-08-23"),
        ("対象システム", "経費精算システム"),
        ("機密区分", "社内一般（架空）"),
        ("作成者", "アプリ設計担当"),
        ("確認者", "業務・基盤担当"),
    ]
    table = document.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(meta_rows):
        set_cell_text(table.cell(index, 0), label, bold=True)
        shade_cell(table.cell(index, 0), "F2F4F7")
        set_cell_text(table.cell(index, 1), value)

    document.add_paragraph()
    purpose = document.add_paragraph(
        "本書は設計書解析・図形認識・表構造抽出・日本語文書処理の検証入力として使用する Word 版の複合設計書"
        "フィクスチャである。内容は経費精算システム_設計書_検証用.xlsx（Excel方眼紙版）と同一の識別子・数値を"
        "共有するが、章構成と表現は Word の慣用的な作り方に合わせて再構成している。"
    )
    purpose.paragraph_format.space_before = Pt(6)

def build_chapter1(document) -> None:
    document.add_heading("第1章　概要", level=1)

    document.add_heading("1.1 文書の目的と設計前提", level=2)
    document.add_paragraph(
        "設計書解析・図形認識・表構造抽出・日本語文書処理の検証入力として使用する。Excel方眼紙の体裁で作成した"
        "原本と同じ識別子・数値を Word 設計書として再構成し、両形式の変換品質を比較できるようにした。"
    )
    p = document.add_paragraph()
    p.add_run("利用者: ").bold = True
    p.add_run("申請者、課長／部長承認者、経理担当者。")
    p2 = document.add_paragraph()
    p2.add_run("認証: ").bold = True
    p2.add_run("OIDC／JWT。Web UI はアクセストークンを API へ送信する。")
    p3 = document.add_paragraph()
    p3.add_run("金額閾値（BR-01）: ").bold = True
    p3.add_run("100,000円以下は課長承認、100,000円超は部長承認とする。")
    p4 = document.add_paragraph()
    p4.add_run("領収書: ").bold = True
    p4.add_run("10,000円以上では添付必須。PDF/JPEG/PNG、最大10MB。")

    document.add_heading("1.2 外部インターフェース一覧", level=2)
    headers = ["I/F ID", "送信元", "送信先", "方式", "Timeout", "Retry", "失敗時の扱い"]
    rows = [
        ["IF-01", "Browser", "Web UI", "HTTPS/TLS1.2+", "30秒", "なし", "画面に再操作案内"],
        ["IF-02", "Web UI", "Expense API", "REST/JSON", "3秒", "1回", "Idempotency-Keyで安全に再送"],
        ["IF-03", "Expense API", "認証基盤", "OIDC/JWT", "2秒", "なし", "401、認証画面へ遷移"],
        ["IF-04", "Expense API", "承認Service", "内部REST", "2秒", "2回", "Outboxへ退避、PENDING維持"],
        ["IF-05", "Outbox Worker", "Event Bus", "AMQP", "5秒", "最大5回", "DLQ、OPS-ALM-02発報"],
        ["IF-06", "Expense API", "会計System", "REST/Batch", "10秒", "最大5回", "ERROR遷移、手動再送"],
    ]
    build_table(document, headers, rows)
    check("D06", "通常表(網掛け+太字ヘッダー行): 外部I/F一覧", "python-docx add_table (native) + raw oxml w:shd shading on header row")

    note = document.add_paragraph()
    note.add_run("責任分界: IF-01〜05は当システム主管、IF-06の接続先仕様・停止時間帯は会計システム主管。調整状況は付録の")
    add_internal_hyperlink(note, "未決事項ISSUE-03", "bm_issue03")
    note.add_run("を参照。")

    document.add_heading("1.3 参考資料", level=2)
    ref = document.add_paragraph("外部連携仕様の詳細は次のAPI仕様書（外部参照）を参照する: ")
    add_external_hyperlink(ref, "経費精算API仕様書", "https://example.invalid/spec/expense-api")
    check(
        "D12", "ハイパーリンク: 外部URL + 文書内ブックマーク相互参照",
        "raw oxml: w:hyperlink r:id=(document.part.relate_to) for external, "
        "w:hyperlink w:anchor= + w:bookmarkStart/End for internal",
    )

def build_chapter2(document) -> None:
    document.add_heading("第2章　状態遷移", level=1)

    document.add_heading("2.1 状態と定義", level=2)
    document.add_paragraph(
        "申請は DRAFT → SUBMITTED → PENDING → APPROVED → PAYMENT_PENDING → PAID の主経路に加え、"
        "RETURNED（差戻し）、ERROR（連携エラー）、CANCELLED（取消）を持つ。状態更新は Expense API のみが実施する。"
    )

    document.add_heading("2.2 状態コード表", level=2)
    status_headers = ["コード", "表示名", "終端", "許可操作"]
    status_rows = [
        ["DRAFT", "下書き", "—", "編集/提出/取消"],
        ["SUBMITTED", "提出済", "—", "参照"],
        ["PENDING", "承認待ち", "—", "承認/差戻し"],
        ["RETURNED", "差戻し", "—", "編集/再提出"],
        ["APPROVED", "承認済", "—", "参照"],
        ["PAYMENT_PENDING", "支払待ち", "—", "参照"],
        ["PAID", "支払済", "○", "参照"],
        ["CANCELLED", "取消", "○", "参照"],
        ["ERROR", "連携エラー", "—", "再送/手動解消"],
    ]
    table = document.add_table(rows=1 + len(status_rows) + 1, cols=4)
    table.style = "Table Grid"
    header_row(table, status_headers)
    for offset, row in enumerate(status_rows, start=1):
        data_row(table, offset, row)
    # D07a: vMerge — PAID and CANCELLED share the same 終端="○" cell vertically.
    # _Cell.merge() appends the second cell's own paragraph text onto the
    # first, so clear one side first or the merged cell reads "○○".
    paid_row = 1 + status_rows.index(["PAID", "支払済", "○", "参照"])
    cancelled_row = 1 + status_rows.index(["CANCELLED", "取消", "○", "参照"])
    table.cell(cancelled_row, 2).text = ""
    merged_terminal_cell = table.cell(paid_row, 2).merge(table.cell(cancelled_row, 2))
    set_cell_text(merged_terminal_cell, "○")
    # D07b: gridSpan — trailing note row spans all four columns.
    note_row_index = 1 + len(status_rows)
    note_cell = table.cell(note_row_index, 0).merge(table.cell(note_row_index, 3))
    set_cell_text(note_cell, "終端状態（○）に達した申請は、いかなる操作によっても他の状態へ遷移しない。")
    check("D07", "gridSpan横結合 + vMerge縦結合を含む表: 状態コード表", "python-docx _Cell.merge() (native; emits w:gridSpan / w:vMerge)")

    document.add_heading("2.3 遷移定義", level=2)
    tr_headers = ["遷移ID", "遷移元", "イベント", "ガード/条件", "遷移先", "副作用"]
    tr_rows = [
        ["TR-01", "—", "新規作成", "認証済み", "DRAFT", "一時保存領域作成"],
        ["TR-02", "DRAFT", "提出", "必須入力OK、金額>0", "SUBMITTED", "監査ログ記録"],
        ["TR-03", "SUBMITTED", "受付", "Idempotency-Key未処理", "PENDING", "承認イベント発行"],
        ["TR-04", "PENDING", "承認", "承認権限あり", "APPROVED", "承認日時を記録"],
        ["TR-05", "PENDING", "差戻し", "理由1文字以上", "RETURNED", "申請者へ通知"],
        ["TR-06", "RETURNED", "再提出", "修正後の入力検証OK", "SUBMITTED", "版番号を加算"],
        ["TR-07", "APPROVED", "会計連携成功", "仕訳番号あり", "PAYMENT_PENDING", "仕訳番号を保存"],
        ["TR-08", "PAYMENT_PENDING", "支払確定", "支払日あり", "PAID", "完了通知"],
        ["TR-09", "PAYMENT_PENDING", "連携失敗", "retry_count >= 5", "ERROR", "DLQ登録・運用アラート"],
        ["TR-10", "ERROR", "手動再送", "原因解消済・運用権限あり", "PAYMENT_PENDING", "retry_countを0へ戻す"],
        ["TR-11", "DRAFT", "取消", "未提出", "CANCELLED", "論理削除"],
    ]
    build_table(document, tr_headers, tr_rows)
    check("D01", "見出し1〜4階層の章立て+本文段落", "python-docx add_heading levels 1-4 (native) across chapters")

    document.add_heading("2.3.1 不変条件", level=3)
    nested_bullets = [
        (0, "終端状態"),
        (1, "PAID"),
        (2, "他状態への遷移不可"),
        (1, "CANCELLED"),
        (2, "再利用不可（DRAFTからの取消のみが遷移元）"),
        (0, "冪等性"),
        (1, "同一Idempotency-Keyの再送"),
        (2, "claim_idを変えず二重登録を防止（DR-04）"),
        (0, "運用権限"),
        (1, "ERRORからの再送（TR-10）"),
        (2, "原因解消済・運用権限ありの場合のみ許可"),
    ]
    bullet_styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
    bullet_num_id = add_multilevel_bullet_numbering(document)
    for level, text in nested_bullets:
        bullet_paragraph = document.add_paragraph(text, style=bullet_styles[level])
        set_explicit_list_level(bullet_paragraph, bullet_num_id, level)
    check("D09", "箇条書き3階層ネスト", "python-docx paragraph style (native) + raw oxml explicit w:numPr/w:ilvl per paragraph (own multi-level abstractNum)")

def build_chapter3(document) -> None:
    document.add_heading("第3章　シーケンス", level=1)

    document.add_heading("3.1 申請受付シーケンス概要", level=2)
    document.add_paragraph(
        "SQ-01。申請者が Web UI で入力・提出すると、Expense API が JWT・入力・権限・冪等性を検証したうえで "
        "PostgreSQL への書き込みを行い、201 応答を返す。Outbox Worker は未送信イベントを取得して Event Bus へ publish し、"
        "承認サービスが at-least-once で approval.requested を受信して承認タスクを生成する。"
    )

    document.add_heading("3.2 トランザクション設計注記", level=2)
    document.add_paragraph(
        "TX-01 は expense_claim テーブルへの INSERT と outbox_event テーブルへの INSERT を同一トランザクションで確定する。"
        "HTTP 応答後のイベント配送は非同期であり、Event Bus 障害時は NFR-OPS-04 の間隔（1/5/30/120/600秒、最大5回）で"
        "再送し、承認サービスは event_id により冪等に処理する。"
    )

    document.add_heading("3.3 エラー応答例", level=2)
    document.add_paragraph("入力検証エラー時は次の ProblemDetails 形式で応答する。")
    code_lines = [
        "HTTP/1.1 400 Bad Request",
        "Content-Type: application/problem+json",
        "",
        "{",
        '  "type": "https://example.invalid/problems/validation-error",',
        '  "title": "入力内容に誤りがあります",',
        '  "status": 400,',
        '  "errors": [{"field": "amount", "code": "EXP-E004"}]',
        "}",
    ]
    code_paragraph = document.add_paragraph(style="Source Code")
    code_run = code_paragraph.add_run(code_lines[0])
    for line in code_lines[1:]:
        code_run.add_break()
        code_run.add_text(line)
    check("D11", "CodeスタイルのHTTP/JSONブロック + インラインcode run", "python-docx custom paragraph style 'Source Code' (native) + run.add_break()/add_text() for embedded newlines")

def build_chapter4(document) -> None:
    document.add_heading("第4章　業務フロー", level=1)

    document.add_heading("4.1 業務ルール・責任分界", level=2)
    br_headers = ["ルールID", "条件/判断", "責任者", "証跡・参照先"]
    br_rows = [
        ["BR-01", "10万円以下は課長、超過は部長承認", "業務主管", "approval_task.route_code"],
        ["BR-02", "差戻し理由は1文字以上500文字以下", "承認者", "audit_log / TR-05"],
        ["BR-03", "会計連携5回失敗でERROR・DLQへ移管", "経理／運用", "TR-09 / OPS-ALM-02"],
        ["SCOPE-01", "銀行振込指示と入金消込は対象外", "会計システム主管", "IF-06責任分界"],
    ]
    build_table(document, br_headers, br_rows)

    document.add_heading("4.2 実施手順", level=2)
    document.add_paragraph("1.", style="List Number").text = "領収書・明細を入力する（BF-01）。"
    document.add_paragraph("2.", style="List Number").text = "形式・必須チェックを行う（BF-02）。"
    document.add_paragraph(
        "上記2手順は Web UI 側で完結し、以降は Expense API 側の処理となる。手順番号はここで途切れずに続く想定である。"
    )
    document.add_paragraph("3.", style="List Number").text = "申請を登録する（BF-03、SUBMITTED、TX-01）。"
    document.add_paragraph("4.", style="List Number").text = "承認者を決定する（BF-04、金額により課長/部長）。"
    check("D10", "番号付きリスト（間に段落を挟んで継続）", "python-docx paragraph style='List Number' (native), interrupted by a normal paragraph")

    document.add_heading("4.3 例外処理", level=3)
    document.add_paragraph(
        "会計連携（BF-07）が失敗した場合は BF-08 再送キューへ登録し、1/5/30/120/600秒の間隔で再送する。"
        "5回失敗すると TR-09 により ERROR へ遷移し、BR-03 に従って DLQ へ移管する。"
    )
    document.add_page_break()
    check("D18", "明示的な改ページ（章間）", "python-docx document.add_page_break() (native; emits w:br w:type=\"page\")")

def add_caption(document, text: str) -> None:
    caption = document.add_paragraph(text, style="Caption")
    return caption

def build_chapter5(document, img01_bytes: bytes, img02_bytes: bytes) -> None:
    document.add_heading("第5章　画面・入力仕様", level=1)

    document.add_heading("5.1 入力項目定義", level=2)
    document.add_paragraph("画面ID: EXP-ENTRY-01　経費申請入力。")
    f_headers = ["項目ID", "項目名", "型/桁", "必須", "検証ルール", "エラーコード"]
    f_rows = [
        ["F-01", "申請日", "date", "○", "当日以前、過去1年以内", "EXP-E001"],
        ["F-02", "部門コード", "char(5)", "○", "部門マスタに存在", "EXP-E002"],
        ["F-03", "摘要", "varchar(100)", "○", "1〜100文字、制御文字不可", "EXP-E003"],
        ["F-04", "金額", "integer", "○", "1〜9,999,999円", "EXP-E004"],
        ["F-05", "税区分", "enum", "○", "10%／8%／非課税", "EXP-E005"],
        ["F-06", "領収書", "file", "条件付", "10,000円以上は必須、最大10MB", "EXP-E006"],
        ["F-07", "備考", "varchar(500)", "—", "500文字以内", "EXP-E007"],
    ]
    build_table(document, f_headers, f_rows)

    document.add_heading("5.2 画面イメージ", level=2)
    document.add_paragraph("文字を含まない画面キャプチャの例を次に示す（IMG-01）。")
    image_paragraph = document.add_paragraph()
    image_run = image_paragraph.add_run()
    image_run.add_picture(io.BytesIO(img01_bytes), width=Emu(3200400))
    add_caption(document, "図 5-1　経費申請入力画面（イメージ、文字なしPNG）")
    check("D13", "画像: inline+キャプション段落 / 浮動(anchor)配置 各1", "python-docx run.add_picture (native, inline) + Caption style; second image converted wp:inline→wp:anchor via raw oxml")

    document.add_heading("5.3 補足資料", level=2)
    document.add_paragraph("入力画面に添える補足資料（日本語文字を含む画像）を、本文の右側に回り込む形で配置する（IMG-02）。")
    image2_paragraph = document.add_paragraph()
    image2_run = image2_paragraph.add_run()
    image2_run.add_picture(io.BytesIO(img02_bytes), width=Emu(2743200))
    convert_last_picture_to_anchor(image2_paragraph)
    add_caption(document, "図 5-2　補足資料（日本語・英数字を含むPNG。OCR検証用）")

    document.add_heading("5.4 エラーメッセージ", level=3)
    msg_headers = ["コード", "表示メッセージ", "表示箇所"]
    msg_rows = [
        ["EXP-E001", "申請日を入力してください。", "申請日直下"],
        ["EXP-E002", "有効な部門コードを入力してください。", "部門コード直下"],
        ["EXP-E003", "摘要は1〜100文字で入力してください。", "摘要直下"],
        ["EXP-E004", "金額は1〜9,999,999円で入力してください。", "金額直下"],
        ["EXP-E006", "10,000円以上の申請には領収書が必要です。", "領収書直下"],
        ["EXP-E409", "この申請は既に受け付けています。", "画面上部"],
    ]
    build_table(document, msg_headers, msg_rows)

def build_chapter6(document) -> None:
    document.add_heading("第6章　API・データ設計", level=1)

    document.add_heading("6.1 REST API一覧", level=2)
    api_headers = ["API ID", "Method", "Path", "目的", "主な応答"]
    api_rows = [
        ["API-01", "POST", "/v1/expense-claims", "申請を提出", "201 / 400 / 401 / 409"],
        ["API-02", "GET", "/v1/expense-claims/{id}", "申請詳細を取得", "200 / 401 / 404"],
        ["API-03", "PATCH", "/v1/expense-claims/{id}", "下書き・差戻しを修正", "200 / 400 / 409"],
        ["API-04", "POST", "/v1/expense-claims/{id}/approve", "承認", "200 / 403 / 409"],
        ["API-05", "POST", "/v1/expense-claims/{id}/return", "差戻し", "200 / 400 / 403"],
        ["API-06", "POST", "/v1/expense-claims/{id}/cancel", "取消", "200 / 409"],
        ["API-07", "GET", "/v1/expense-claims", "一覧・検索", "200 / 400 / 401"],
    ]
    build_table(document, api_headers, api_rows)

    document.add_heading("6.1.1 リクエストヘッダ", level=3)
    document.add_paragraph(
        "API-01 はヘッダーに Authorization（Bearer JWT、必須）と Idempotency-Key（uuid、必須）を要求する。"
    )
    document.add_heading("6.1.1.1 Idempotency-Key の扱い", level=4)
    document.add_paragraph(
        "Idempotency-Key は同一利用者について24時間一意として扱う。同一キーでの再送は新規登録を行わず、"
        "既存の claim_id と同じ結果を返す（DR-04）。"
    )

    document.add_heading("6.2 主要テーブルとイベント", level=2)
    outer = document.add_table(rows=2, cols=5)
    outer.style = "Table Grid"
    header_row(outer, ["区分", "名称", "型", "必須", "説明"])
    data_row(outer, 1, ["Body", "amount", "integer", "○", ""])
    detail_cell = outer.cell(1, 4)
    detail_cell.text = ""
    detail_cell.paragraphs[0].add_run("円、1〜9,999,999。税込/税抜の内訳は下表のとおり。")
    nested = detail_cell.add_table(rows=2, cols=2)
    nested.style = "Table Grid"
    set_cell_text(nested.cell(0, 0), "税込金額", bold=True)
    set_cell_text(nested.cell(0, 1), "12,800円")
    set_cell_text(nested.cell(1, 0), "税抜金額(10%)", bold=True)
    set_cell_text(nested.cell(1, 1), "11,637円")
    check("D08", "ネスト表1箇所", "python-docx _Cell.add_table() (native)")

    outbox = document.add_paragraph()
    outbox.add_run("outbox_event.status は ")
    status_run = outbox.add_run("outbox_event.status")
    status_run.font.name = "Consolas"
    outbox.add_run(
        " の値として NEW / SENDING / SENT / FAILED のいずれかを取り、retry_count の上限は5である。"
        "上限に達すると TR-09 に従い ERROR へ遷移し DLQ に登録される。"
    )

def build_chapter7(document) -> None:
    document.add_heading("第7章　テストシナリオ", level=1)
    document.add_heading("7.1 代表シナリオ", level=2)
    tc_headers = ["TC-ID", "観点/入力", "通過経路", "期待HTTP", "期待最終状態", "関連設計"]
    tc_rows = [
        ["TC-001", "12,800円、領収書あり", "課長承認→会計成功", "201", "PAID", "TR-02/03/04/07/08, API-01"],
        ["TC-002", "150,000円、領収書あり", "部長承認→会計成功", "201", "PAID", "シーケンス alt 高額経路"],
        ["TC-003", "20,000円、領収書なし", "入力検証NG", "400", "DRAFT", "F-06, EXP-E006"],
        ["TC-004", "同一Idempotency-Keyを再送", "既存結果を返却", "201", "SUBMITTED", "DR-04, API-01"],
        ["TC-005", "承認者が理由付き差戻し", "PENDING→RETURNED", "200", "RETURNED", "TR-05"],
        ["TC-006", "会計I/Fが5xxを返す", "再送キュー→成功", "200", "PAYMENT_PENDING", "業務フロー 連携失敗"],
        ["TC-007", "支払済申請を更新", "不変条件違反", "409", "PAID", "DR-09"],
        ["TC-008", "Event Busを5分停止", "Outbox再送→承認タスク生成", "201", "PENDING", "SQ-01, EVT-01, NFR-OPS-04"],
        ["TC-009", "会計I/Fが5回連続失敗", "DLQ→ERROR→手動再送", "200", "ERROR", "TR-09/10, BR-03, OPS-ALM-02"],
    ]
    build_table(document, tc_headers, tc_rows)

def build_chapter8(document):
    document.add_heading("第8章　非機能・運用設計", level=1)
    document.add_paragraph(
        "IPA非機能要求グレードの分類を参考に、目標値・測定方法・実装・運用判断を追跡可能にする。次節の一覧は"
        "列数が多いため、本節のみ横向きページで示す。"
    )

    portrait_width, portrait_height = document.sections[-1].page_width, document.sections[-1].page_height
    landscape_section = document.add_section(WD_SECTION.NEW_PAGE)
    landscape_section.orientation = WD_ORIENT.LANDSCAPE
    landscape_section.page_width, landscape_section.page_height = portrait_height, portrait_width
    landscape_section.header.is_linked_to_previous = False
    landscape_section.header.paragraphs[0].text = "経費精算システム 基本・詳細設計書（8.1 非機能要求一覧・横向き）"

    document.add_heading("8.1 非機能要求・SLO", level=2)
    nfr_headers = ["NFR ID", "分類", "目標値/条件", "測定・判定", "設計上の実現手段", "合意状態"]
    nfr_rows = [
        ["NFR-AVL-01", "可用性", "平日8:00–22:00の月間稼働率99.9%以上", "外形監視1分間隔。計画停止を除外", "2AZ配置、ヘルスチェック、ローリング更新", "条件付合意"],
        ["NFR-AVL-02", "回復性", "RTO 4時間／RPO 15分", "年1回の復旧訓練で実測", "日次バックアップ＋WAL 15分転送", "要確認"],
        ["NFR-PERF-01", "性能", "通常申請 p95 1.0秒以下、20 TPS", "APMで5分窓。添付転送時間を除外", "DB索引、接続プール、非同期イベント化", "合意済"],
        ["NFR-CAP-01", "拡張性", "月5万申請、3年分をオンライン保持", "月次容量レポート、70%で警告", "claim_id分散、領収書はObject Storage", "要確認"],
        ["NFR-SEC-01", "認証認可", "承認APIは承認者ロールのみ。管理操作はMFA", "403率・権限テスト・監査ログ", "OIDC、RBAC、最小権限", "合意済"],
        ["NFR-OPS-04", "再送", "1/5/30/120/600秒、最大5回", "retry_countとnext_attempt_atを照合", "Outbox Worker、指数バックオフ、DLQ", "合意済"],
    ]
    build_table(document, nfr_headers, nfr_rows)
    check("D05", "セクション区切り: 縦→横向き(ワイド表)→縦。第2セクションは別ヘッダー", "python-docx document.add_section + WD_ORIENT.LANDSCAPE (native) + header.is_linked_to_previous=False")

    portrait_section = document.add_section(WD_SECTION.NEW_PAGE)
    portrait_section.orientation = WD_ORIENT.PORTRAIT
    portrait_section.page_width, portrait_section.page_height = portrait_width, portrait_height
    portrait_section.header.is_linked_to_previous = False
    portrait_section.header.paragraphs[0].text = "経費精算システム 基本・詳細設計書"

    document.add_heading("8.2 監視・障害対応マトリクス", level=2)
    ops_headers = ["Alarm ID", "検知条件", "自動処理", "エスカレーション", "SLA"]
    ops_rows = [
        ["OPS-ALM-01", "API 5xx率 > 5% / 5分", "新規Pod追加、詳細ログ採取", "15分継続で開発責任者", "5分"],
        ["OPS-ALM-02", "DLQ件数 >= 1", "対象event_idを凍結", "30分で経理・基盤へ連絡", "即時"],
        ["OPS-ALM-04", "会計I/F timeout 3回 / 5分", "サーキットをopen、再送へ退避", "15分で経理主管", "5分"],
    ]
    build_table(document, ops_headers, ops_rows)

    document.add_heading("8.3 設計判断・未決事項", level=2)
    add_bookmark(document.paragraphs[-1], "bm_issue03")
    issue_headers = ["ID", "論点", "現時点の判断/選択肢", "Owner", "期限", "Status"]
    issue_rows = [
        ["ADR-001", "イベント原子性", "Transactional Outboxを採用", "アプリTL", "2026-08-23", "採用"],
        ["ADR-002", "承認連携方式", "同期RESTではなくEvent Bus経由", "基盤TL", "2026-08-23", "採用"],
        ["ISSUE-01", "RPO 15分の費用", "WAL転送費と業務影響を比較", "基盤担当", "2026-08-28", "要確認"],
        ["ISSUE-02", "3年超データ", "アーカイブ先・検索SLAを業務合意", "業務担当", "2026-08-30", "要確認"],
        ["ISSUE-03", "会計停止時間帯", "IF-06再送窓を会計主管と調整", "経理担当", "2026-08-27", "対応中"],
    ]
    build_table(document, issue_headers, issue_rows)
    return "bm_issue03"

def build_appendix(document, notes: NoteBook, landscape_bookmark: str) -> None:
    document.add_heading("付録", level=1)

    document.add_heading("付録A　表記ルールの補足", level=2)
    document.add_paragraph(
        "図形の意味を先に定義し、読み手による解釈差を抑える。識別子は状態TR、API、項目F、データDR、テストTC、"
        "非機能NFR、業務ルールBR、外部I/F IFで相互参照する。"
    )
    document.add_heading("付録A.1　図形凡例の補足", level=7)
    check("D02", "Heading 7を付録に1箇所", "python-docx add_heading(level=7) (native; style 'Heading 7')")
    document.add_paragraph(
        "本節は見出しレベル7で登録している。HeadingLevel の正規表現が1〜6のみを許可する実装では、"
        "この段落は通常段落として扱われる想定である。"
    )

    document.add_heading("付録B　用語・識別子対照表", level=2)
    glossary_headers = ["識別子", "種別", "意味"]
    glossary_rows = [
        ["TR", "状態遷移", "02_振る舞い図の遷移定義（TR-01〜TR-11）"],
        ["API", "REST API", "06_API・データのエンドポイント定義（API-01〜API-07）"],
        ["F", "画面入力項目", "05_画面・入力仕様の項目定義（F-01〜F-07）"],
        ["DR", "データ整合性ルール", "06_API・データの整合性・監査ルール（DR-01〜DR-09）"],
        ["TC", "テストケース", "07_テストシナリオの代表ケース（TC-001〜TC-009）"],
        ["NFR", "非機能要求", "08_非機能・運用のSLO一覧"],
        ["BR", "業務ルール", "04_業務フローの責任分界"],
        ["IF", "外部インターフェース", "01_システム概要の外部I/F一覧（IF-01〜IF-06）"],
    ]
    build_table(document, glossary_headers, glossary_rows)

    document.add_heading("付録C　レビュー注記", level=2)

    revision = document.add_paragraph()
    revision.add_run("初版レビューでは、")
    add_tracked_deletion(revision, "承認者は課長のみとする方針であったが、")
    add_tracked_insertion(revision, "金額に応じて課長または部長が承認する方式（BR-01）に見直した。")
    revision.add_run("本改版で反映済み。")
    check("D17", "コメント1件 + 未承認の変更履歴（挿入/削除）1組", "docx.Document.add_comment (native) + raw oxml w:ins/w:del")

    comment_target = document.add_paragraph()
    comment_target.add_run(
        "8.3節のISSUE-03（会計停止時間帯）は経理主管との調整が完了し次第、本書を1.3版へ改版する。詳細は"
    )
    ref = comment_target.add_run("8.3 設計判断・未決事項")
    comment_target.add_run("を参照。")
    document.add_comment(ref, text="経理主管との調整完了予定日を明記してください。", author=REVIEW_AUTHOR, initials="RV")

    footnote_p = document.add_paragraph("本書は検証用サンプルである。")
    notes.add_footnote(footnote_p, "実在のシステム・組織・金額を示すものではなく、DRMD変換検証のために作成した架空の設計内容である。")

    endnote_p = document.add_paragraph("調査・参考資料の詳細な参照条件は文末脚注を参照。")
    notes.add_endnote(endnote_p, "参照URLはいずれも2026-08-23時点の調査結果であり、実アクセス確認は行っていない（検証用ダミー記載を含む）。")

    callout_lead = document.add_paragraph("本書の位置づけについて、次のコールアウトを参照。")
    add_textbox(callout_lead, "注記: 本書はDRMD DOCX変換検証用のフィクスチャであり、Word設計書の一般的な要素を意図的に組み合わせている。")
    document.add_paragraph()
    check("D14", "テキストボックス（コールアウト注記）", "raw oxml legacy VML: w:pict > v:shape > v:textbox > w:txbxContent")

    deco = document.add_paragraph()
    deco.add_run("書式サンプル: ")
    bold_run = deco.add_run("太字")
    bold_run.bold = True
    deco.add_run(" / ")
    italic_run = deco.add_run("斜体")
    italic_run.italic = True
    deco.add_run(" / ")
    underline_run = deco.add_run("下線")
    underline_run.underline = True
    deco.add_run(" / ")
    strike_run = deco.add_run("取消線")
    strike_run.font.strike = True
    deco.add_run(" / ")
    color_run = deco.add_run("強調色")
    color_run.font.color.rgb = RGBColor.from_string(RED_HEX)
    deco.add_run(" / ")
    highlight_run = deco.add_run("ハイライト")
    highlight_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    check("D15", "run装飾: 太字/斜体/下線/取消線/文字色/ハイライト", "python-docx Run.bold/italic/underline/font.strike/font.color.rgb/font.highlight_color (native)")

# --------------------------------------------------------------------------
# Self-verification: reopen the generated file and confirm each element.
# --------------------------------------------------------------------------

def verify(output_path: Path) -> dict[str, bool]:
    verified: dict[str, bool] = {}

    reopened = docx.Document(str(output_path))
    heading_styles = {p.style.name for p in reopened.paragraphs if p.style and p.style.name.startswith("Heading")}
    verified["D01"] = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}.issubset(heading_styles)
    verified["D02"] = "Heading 7" in heading_styles
    verified["D06"] = len(reopened.tables) >= 1
    verified["D09"] = any(p.style and p.style.name in ("List Bullet", "List Bullet 2", "List Bullet 3") for p in reopened.paragraphs)
    verified["D10"] = any(p.style and p.style.name == "List Number" for p in reopened.paragraphs)
    verified["D13"] = len(reopened.inline_shapes) >= 1
    verified["D17_comment"] = len(reopened.comments) >= 1

    with zipfile.ZipFile(output_path) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8")
        content_types = archive.read("[Content_Types].xml").decode("utf-8")

        footer_xml = "".join(
            archive.read(name).decode("utf-8") for name in names if name.startswith("word/footer") and name.endswith(".xml")
        )
        header_xml = "".join(
            archive.read(name).decode("utf-8") for name in names if name.startswith("word/header") and name.endswith(".xml")
        )

        verified["D03"] = " TOC " in document_xml and "fldCharType=\"separate\"" in document_xml
        verified["D04"] = "fldSimple" in footer_xml and "w:instr=\"PAGE\"" in footer_xml and "経費精算" in header_xml
        verified["D05_sections"] = document_xml.count("<w:sectPr") >= 3 or document_xml.count("<w:sectPr>") >= 1
        verified["D05_landscape"] = "w:orient=\"landscape\"" in document_xml
        verified["D07_gridspan"] = "gridSpan" in document_xml
        verified["D07_vmerge"] = "vMerge" in document_xml
        verified["D09_ilvl"] = 'w:ilvl w:val="1"' in document_xml and 'w:ilvl w:val="2"' in document_xml
        verified["D08_nested_tables"] = document_xml.count("<w:tbl>") > len(reopened.tables)
        verified["D11"] = "SourceCode" in document_xml or "Source Code" in document_xml
        verified["D12_external"] = "Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink\"" in archive.read("word/_rels/document.xml.rels").decode("utf-8")
        verified["D12_internal"] = "w:anchor=\"bm_issue03\"" in document_xml and "w:bookmarkStart" in document_xml
        verified["D13_anchor"] = "<wp:anchor" in document_xml
        verified["D14"] = "txbxContent" in document_xml and "v:shape" in document_xml
        verified["D15_color_highlight"] = "w:highlight" in document_xml and "<w:color w:val=" in document_xml
        verified["D16_footnote_part"] = "word/footnotes.xml" in names and "footnoteReference" in document_xml
        verified["D16_endnote_part"] = "word/endnotes.xml" in names and "endnoteReference" in document_xml
        verified["D17_tracked"] = "<w:ins " in document_xml and "<w:del " in document_xml
        verified["D18_pagebreak"] = "w:type=\"page\"" in document_xml
        verified["ContentTypes_footnotes"] = "footnotes.xml" in content_types
        verified["ContentTypes_endnotes"] = "endnotes.xml" in content_types
        verified["Q12_no_leak"] = "OCR-JP-20260823-017" not in document_xml

    return verified

def print_checklist(verified: dict[str, bool]) -> bool:
    print()
    print("=" * 78)
    print("complex-design-doc.docx element checklist (D01-D18)")
    print("=" * 78)
    all_ok = True
    seen_ids = []
    for item_id, label, method, _note in CHECKLIST:
        if item_id in seen_ids:
            continue
        seen_ids.append(item_id)
        related = {k: v for k, v in verified.items() if k == item_id or k.startswith(item_id + "_")}
        if related:
            ok = all(related.values())
        else:
            ok = True  # no independent probe defined; construction succeeded without raising
        all_ok = all_ok and ok
        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {item_id:4s} {label}")
        print(f"       method: {method}")
        if related:
            detail = ", ".join(f"{k}={v}" for k, v in related.items())
            print(f"       probe:  {detail}")
    print("-" * 78)
    print("Additional structural probes:")
    for key in ("ContentTypes_footnotes", "ContentTypes_endnotes", "Q12_no_leak"):
        if key in verified:
            print(f"  {key}: {verified[key]}")
    print("=" * 78)
    return all_ok

# --------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------

def main(argv=None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output", nargs="?", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--xlsx", type=Path, default=DEFAULT_XLSX)
    parser.add_argument("--root", type=Path, default=DEFAULT_ROOT)
    args = parser.parse_args(argv)

    if not args.xlsx.exists():
        print(f"error: source xlsx not found: {args.xlsx}", file=sys.stderr)
        return 2

    output_path = args.output
    output_path.parent.mkdir(parents=True, exist_ok=True)
    build(output_path, args.xlsx)
    print(f"wrote {output_path} ({output_path.stat().st_size} bytes)")

    verified = verify(output_path)
    ok = print_checklist(verified)
    return 0 if ok else 1

if __name__ == "__main__":
    raise SystemExit(main())